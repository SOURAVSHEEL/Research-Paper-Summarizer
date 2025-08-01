import PyPDF2
import docx
import markdown
import streamlit as st
from utils.logger import log_info, log_error, log_warning

def extract_text_from_document(uploaded_file):
    """Extract text from various document formats"""
    file_type = uploaded_file.name.lower().split('.')[-1]
    
    log_info(f"Processing {file_type.upper()} file: {uploaded_file.name}")
    
    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_type in ['doc', 'docx']:
            return extract_text_from_docx(uploaded_file)
        elif file_type == 'txt':
            return extract_text_from_txt(uploaded_file)
        elif file_type in ['md', 'markdown']:
            return extract_text_from_markdown(uploaded_file)
        else:
            log_error(f"Unsupported file type: {file_type}")
            return None
    except Exception as e:
        log_error(f"Failed to extract text from {uploaded_file.name}: {str(e)}")
        return None

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        log_info("Extracting text from PDF")
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text
            log_info(f"Processed PDF page {i+1}/{len(pdf_reader.pages)}")
        
        log_info(f"Successfully extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        log_error(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    try:
        log_info("Extracting text from DOCX")
        doc = docx.Document(docx_file)
        text = ""
        
        # Extract text from paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
            if i % 50 == 0:  # Log progress every 50 paragraphs
                log_info(f"Processed {i+1}/{len(doc.paragraphs)} paragraphs")
        
        # Extract text from tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if table_count > 0:
            log_info(f"Extracted text from {table_count} tables")
        
        log_info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        log_error(f"DOCX extraction error: {str(e)}")
        return None

def extract_text_from_txt(txt_file):
    """Extract text from TXT file"""
    try:
        log_info("Extracting text from TXT")
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = txt_file.read().decode(encoding)
                log_info(f"Successfully decoded TXT with {encoding} encoding")
                log_info(f"Extracted {len(text)} characters from TXT")
                return text
            except UnicodeDecodeError:
                continue
        
        log_error("Failed to decode TXT file with any supported encoding")
        return None
    except Exception as e:
        log_error(f"TXT extraction error: {str(e)}")
        return None

def extract_text_from_markdown(md_file):
    """Extract text from Markdown file"""
    try:
        log_info("Extracting text from Markdown")
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                md_content = md_file.read().decode(encoding)
                log_info(f"Successfully decoded Markdown with {encoding} encoding")
                
                # Convert markdown to plain text (remove markdown syntax)
                html = markdown.markdown(md_content)
                # Remove HTML tags for plain text
                import re
                text = re.sub('<[^<]+?>', '', html)
                # Clean up extra whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)
                
                log_info(f"Extracted {len(text)} characters from Markdown")
                return text
            except UnicodeDecodeError:
                continue
        
        log_error("Failed to decode Markdown file with any supported encoding")
        return None
    except Exception as e:
        log_error(f"Markdown extraction error: {str(e)}")
        return None

def get_document_info(uploaded_file):
    """Get document metadata and information"""
    file_type = uploaded_file.name.lower().split('.')[-1]
    
    info = {
        'name': uploaded_file.name,
        'size': uploaded_file.size,
        'type': file_type.upper(),
        'pages': 0
    }
    
    try:
        if file_type == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            info['pages'] = len(pdf_reader.pages)
            if pdf_reader.metadata:
                info['title'] = pdf_reader.metadata.get('/Title', 'Unknown')
                info['author'] = pdf_reader.metadata.get('/Author', 'Unknown')
        elif file_type in ['doc', 'docx']:
            doc = docx.Document(uploaded_file)
            info['pages'] = len(doc.paragraphs) // 20  # Rough estimate
            core_props = doc.core_properties
            info['title'] = core_props.title or 'Unknown'
            info['author'] = core_props.author or 'Unknown'
        
        log_info(f"Document info extracted: {info}")
        return info
    except Exception as e:
        log_warning(f"Could not extract document info: {str(e)}")
        return info

def is_supported_file_type(filename):
    """Check if file type is supported"""
    supported_extensions = ['pdf', 'doc', 'docx', 'txt', 'md', 'markdown']
    file_extension = filename.lower().split('.')[-1]
    return file_extension in supported_extensions
